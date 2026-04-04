"""
PDF Generator Service for Music License Documents

Handles dynamic PDF generation using PDF -> DOCX -> Edit -> PDF workflow.
"""

import os
import sys
import platform
import subprocess
from typing import Dict, Optional
from datetime import datetime
import hashlib
from pdf2docx import Converter
from docx import Document

class PDFGenerator:
    """
    Generates secure license PDFs by converting to DOCX, editing, and converting back.
    """
    
    def __init__(self):
        self.templates_dir = os.path.join(os.path.dirname(__file__), '..', 'pdf_templates')

    def generate_checksum(self, pdf_bytes: bytes) -> str:
        """
        Generate SHA-256 checksum of the PDF.
        
        Args:
            pdf_bytes: PDF content as bytes
            
        Returns:
            str: SHA-256 checksum in hexadecimal
        """
        sha256_hash = hashlib.sha256()
        sha256_hash.update(pdf_bytes)
        return sha256_hash.hexdigest()

        
    def generate_license_pdf(
        self,
        license_type: str,
        license_data: Dict,
        form_data: Dict,
        transaction_data: Dict
    ) -> Optional[bytes]:
        """
        Generate a complete license PDF.
        
        Args:
            license_type: Type of license ('leasing', 'exclusive', 'unlimited')
            license_data: License and music information
            form_data: User-filled form data
            transaction_data: Payment and transaction details
            
        Returns:
            bytes: Complete PDF as bytes, or None if failed
        """
        # Load the appropriate template
        template_path = self._get_template_path(license_type)
        print(f"[PDF] Template resolved for '{license_type}': {template_path}")
        
        # Prepare replacements
        replacements = self._get_replacements(license_data, form_data, transaction_data)
        
        # Define temp file paths
        # Use a unique ID to avoid collisions if running concurrently (though this is likely single threaded per request)
        import uuid
        run_id = str(uuid.uuid4())[:8]
        temp_docx = os.path.join(os.path.dirname(template_path), f"temp_{run_id}.docx")
        output_pdf = os.path.join(os.path.dirname(template_path), f"output_{run_id}.pdf")
        
        try:
            print("\n" + "="*60)
            print("PDF CONVERSION AND EDITING WORKFLOW")
            print("="*60)
            
            # Step 1: PDF to DOCX
            print("\n[1/3] Converting PDF to DOCX...")
            if not self._pdf_to_docx(template_path, temp_docx):
                return None
            
            # Step 2: Edit DOCX
            print("\n[2/3] Editing DOCX...")
            if not self._edit_docx(temp_docx, replacements):
                return None
            
            # Step 3: DOCX to PDF
            print("\n[3/3] Converting DOCX back to PDF...")
            if not self._docx_to_pdf_libreoffice(temp_docx, output_pdf):
                return None
            
            print("\n" + "="*60)
            print(f"✓ SUCCESS! Created: {output_pdf}")
            print("="*60 + "\n")
            
            # Read the generated PDF bytes
            with open(output_pdf, 'rb') as f:
                pdf_bytes = f.read()
                
            return pdf_bytes
            
        except Exception as e:
            print(f"\n✗ Unexpected error: {e}")
            return None
        finally:
            # Cleanup temp files
            for f in [temp_docx, output_pdf]:
                if os.path.exists(f):
                    try:
                        os.remove(f)
                        print(f"Cleaned up temporary file: {f}")
                    except:
                        pass

    def _get_template_path(self, license_type: str) -> str:
        """Get the path to the appropriate template PDF."""
        templates = {
            'leasing': 'leasing_license_2025.pdf',
            'exclusive': 'exclusive_license_2025.pdf',
            'unlimited': 'unlimited_license_2025.pdf'
        }

        # Normalize license type input
        lt_input = license_type
        lt_raw = getattr(license_type, 'value', license_type)
        lt = str(lt_raw).strip().lower()
        alias_map = {
            'l': 'leasing',
            'lease': 'leasing',
            'leasing': 'leasing',
            'exclusive': 'exclusive',
            'e': 'exclusive',
            'unlimited': 'unlimited',
            'u': 'unlimited',
        }
        lt = alias_map.get(lt, lt)
        
        filename = templates.get(lt)
        if not filename:
            raise ValueError(f"Invalid license type: {license_type}")
            
        path = os.path.join(self.templates_dir, filename)
        if not os.path.exists(path):
            raise FileNotFoundError(f"Template not found: {path}")
            
        return path

    def _get_replacements(self, license_data: Dict, form_data: Dict, transaction_data: Dict) -> Dict[str, str]:
        """Generate the replacement dictionary."""
        
        # Helper to safely get values
        def get_val(data, key, default='N/A'):
            return str(data.get(key) or default)

        # Construct the mapping based on what was in PDFPlaceholderReplacer
        # and what is likely in the document
        
        replacements = {
            '[Beat Name]': get_val(license_data, 'music_name'),
            '[ORDER ID]': get_val(transaction_data, 'order_id'),
            '[BB-L-YYYY-XXXXX]': get_val(transaction_data, 'license_id_formatted'),
            '[BB-E-YYYY-XXXXX]': get_val(transaction_data, 'license_id_formatted'),
            '[BB-U-YYYY-XXXXX]': get_val(transaction_data, 'license_id_formatted'),
            '[Author Full Legal Name]': get_val(form_data, 'author_legal_name') or get_val(form_data, 'author_name'),
            '[Author IPI/CAE Number]': get_val(form_data, 'pro_ipi_number') or get_val(form_data, 'buyer_ip'),
            '[Author PRO]': get_val(form_data, 'pro_name') or get_val(form_data, 'author_name'),
            '[ Author IPI ]': get_val(form_data, 'buyer_ip'),
            '[ Author Full legal Name and Stage Name ]': get_val(form_data, 'author_legal_name') or get_val(form_data, 'author_name'),
            '[ Publisher Name ]': get_val(form_data, 'publisher_name', 'Self-Published'),
            '[Licensee\'s PRO]': get_val(form_data, 'pro_name'),
            '[PRICE_PAID]': f"{float(transaction_data.get('amount_paid', 0)):.2f}",
            '[Timestamp]': get_val(form_data, 'signed_at', datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')),
            '[User IP]': get_val(transaction_data, 'buyer_ip'),
            '[Buyer Full Name]': get_val(form_data, 'licensee_name'),
            '[Image Signature]': 'Bujar Malaj', # Updated to match printed name
            '[ Publisher IPI ]': 'N/A',
            '[Licensee\'s Publisher IPI]': 'N/A',
            '[ automatic enter the [Beat Name]': get_val(license_data, 'music_name'),
            
            # Add combined fields if needed
            '[ Author Full legal  \nName and Stage \nName ]': f"{get_val(form_data, 'author_legal_name')}\n{get_val(form_data, 'artist_stage_name')}",
        }
        
        return replacements

    def _pdf_to_docx(self, pdf_path, docx_path):
        """Convert PDF to DOCX"""
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
            print(f"✓ Converted {pdf_path} to {docx_path}")
            return True
        except Exception as e:
            print(f"✗ Error converting PDF to DOCX: {e}")
            return False

    def _edit_docx(self, docx_path, replacements):
        """
        Edit DOCX file by replacing text, handling split runs.
        """
        try:
            doc = Document(docx_path)
            changes_made = 0
            
            def replace_in_paragraph(paragraph):
                nonlocal changes_made
                # Fast check
                if not any(key in paragraph.text for key in replacements):
                    return

                # 1. Try simple run replacement first (preserves formatting best)
                for run in paragraph.runs:
                    for key, value in replacements.items():
                        if key in run.text:
                            run.text = run.text.replace(key, value)
                            changes_made += 1
                
                # 2. Check if placeholders remain (meaning they are split across runs)
                full_text = paragraph.text
                if any(key in full_text for key in replacements):
                    # We have split placeholders. We must consolidate runs.
                    # This preserves the style of the FIRST run.
                    
                    # Capture style from first run
                    style_data = {}
                    if paragraph.runs:
                        r0 = paragraph.runs[0]
                        style_data = {
                            'bold': r0.bold,
                            'italic': r0.italic,
                            'underline': r0.underline,
                            'font_name': r0.font.name,
                            'font_size': r0.font.size,
                            'color': r0.font.color.rgb if r0.font.color else None
                        }
                    
                    # Perform replacements on the full text
                    for key, value in replacements.items():
                        if key in full_text:
                            full_text = full_text.replace(key, value)
                            changes_made += 1
                    
                    # Rebuild paragraph
                    paragraph.clear()
                    new_run = paragraph.add_run(full_text)
                    
                    # Restore style
                    new_run.bold = style_data.get('bold')
                    new_run.italic = style_data.get('italic')
                    new_run.underline = style_data.get('underline')
                    if style_data.get('font_name'):
                        new_run.font.name = style_data.get('font_name')
                    if style_data.get('font_size'):
                        new_run.font.size = style_data.get('font_size')
                    if style_data.get('color'):
                        new_run.font.color.rgb = style_data.get('color')

            # Edit paragraphs in body
            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph)
                
                # Force add email if missing in footer (User Request)
                if "Jurisdiction: Kosovo" in paragraph.text and "info@bujaabeats.com" not in paragraph.text:
                    paragraph.add_run("\nEmail: info@bujaabeats.com")
                    changes_made += 1
                    print("✓ Forcefully added missing email to footer")
            
            # Edit tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)
            
            doc.save(docx_path)
            print(f"✓ Edited {docx_path} ({changes_made} replacements made)")
            return True
        except Exception as e:
            print(f"✗ Error editing DOCX: {e}")
            return False

    def _find_libreoffice(self):
        """Find LibreOffice installation path"""
        system = platform.system()
        
        if system == "Windows":
            possible_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                r"C:\Program Files\LibreOffice 7\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice 7\program\soffice.exe",
                # Add common paths or check registry if needed
            ]
            for path in possible_paths:
                if os.path.exists(path):
                    return path
            return None
        
        elif system == "Linux":
            # Check if libreoffice is in PATH
            try:
                result = subprocess.run(['which', 'libreoffice'], 
                                      capture_output=True, text=True)
                if result.returncode == 0:
                    return "libreoffice"
            except:
                pass
            return None
        
        elif system == "Darwin":  # macOS
            path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
            return path if os.path.exists(path) else None
        
        return None

    def _docx_to_pdf_libreoffice(self, docx_path, output_pdf_path):
        """Convert DOCX to PDF using LibreOffice"""
        libreoffice_path = self._find_libreoffice()
        
        if not libreoffice_path:
            print("✗ LibreOffice not found!")
            print("\nPlease install LibreOffice:")
            print("  Windows: https://www.libreoffice.org/download/download/")
            print("  Linux: sudo apt-get install libreoffice-writer")
            print("  macOS: brew install --cask libreoffice")
            return False
        
        print(f"Using LibreOffice at: {libreoffice_path}")
        
        output_dir = os.path.dirname(os.path.abspath(output_pdf_path))
        if not output_dir:
            output_dir = os.getcwd()
        
        try:
            # Run LibreOffice conversion
            # Note: LibreOffice output filename is determined by input filename
            result = subprocess.run([
                libreoffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                os.path.abspath(docx_path)
            ], capture_output=True, text=True, timeout=60)
            
            if result.returncode != 0:
                print(f"✗ LibreOffice conversion failed!")
                print(f"stderr: {result.stderr}")
                return False
            
            # LibreOffice creates PDF with same base name as DOCX
            base_name = os.path.splitext(os.path.basename(docx_path))[0]
            temp_pdf = os.path.join(output_dir, base_name + '.pdf')
            
            # Rename if needed
            if temp_pdf != output_pdf_path:
                if os.path.exists(output_pdf_path):
                    os.remove(output_pdf_path)
                if os.path.exists(temp_pdf):
                    os.rename(temp_pdf, output_pdf_path)
                else:
                    print(f"✗ Expected output file {temp_pdf} not found")
                    return False
            
            print(f"✓ Converted {docx_path} to {output_pdf_path}")
            return True
            
        except subprocess.TimeoutExpired:
            print("✗ Conversion timed out (file too large?)")
            return False
        except Exception as e:
            print(f"✗ Error during conversion: {e}")
            return False
